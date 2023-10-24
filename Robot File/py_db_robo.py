import pandas as pd
import numpy as np
import pyodbc as od
import mysql.connector as db
from docx.shared import Inches
import docx
import os
import shutil
from dotenv import load_dotenv
from dbconnection import dbconnection
import pyautogui
from PIL import Image

mylist=['Screen1','Screen2','Screen3']
mylist1=[]

class py_db_robo:
    load_dotenv()
    def Db_Connection(self):
        mysqldbconnection=dbconnection()
        connection_string=mysqldbconnection.dbconnect_string('conn_string')
        conn=db.connect(**connection_string)
        cursor=conn.cursor()
        query_ip=f"SELECT * FROM py_automate where Run_Indicator='Y';"
        cursor.execute(query_ip)
        res = cursor.fetchall()
        record = pd.DataFrame(res)
        Scenario,F_Name,L_Name,Email,Comment = ([] for i in range(5))
        for i in range (len(record)):
            Scenario.append(record[0][i])
            F_Name.append(record[2][i])
            L_Name.append(record[3][i])
            Email.append(record[4][i])
            Comment.append(record[5][i])
        mydict={"SCENARIO":Scenario,"F_NAME":F_Name,"L_NAME":L_Name,"EMAIL":Email,"COMMENT":Comment}
        cursor.close()
        conn.close()
        directory = "temp"
        parent_dir = "D:\End to End\Evidence"
        path = os.path.join(parent_dir,directory)
        if os.path.isdir(path)==False:
                os.mkdir(path)
        return mydict
    
    def DB_Update(self,value,data):
        mysqldbconnection=dbconnection()
        connection_string=mysqldbconnection.dbconnect_string('conn_string')
        conn=db.connect(**connection_string)
        cursor=conn.cursor()
        cursor.execute("UPDATE py_automate SET Result='"+value+"' WHERE Scenario='"+data+"' and Run_Indicator='Y';")  
        conn.commit()
        cursor.close()
        conn.close()
     
     
    def Screenshot_alter(self,SS_msg):
        try:
            mylist.append(SS_msg)
            directory = "temp"
            parent_dir = "D:\End to End\Evidence"
            path = os.path.join(parent_dir,directory)
            if os.path.isdir(path)==False:
                os.mkdir(path)
            #driver.save_screenshot('D:\End to End\Evidence\\'+directory+'\\'+SS_msg+'.png')
        except:
            pass

    def word_cov(self,scn):
        dirr= "temp"
        doc = docx.Document()
        doc.add_heading('Evidence', 0)
        dummy=len(mylist)
        for i in range(dummy):
            doc.add_heading(mylist[i], 5)
            doc.add_picture('D:\End to End\Evidence\\'+dirr+'\\'+mylist[i]+'.png', width=Inches(6.5), height=Inches(3.5))
        doc.save('D:\End to End\Evidence\\'+scn+'.docx')
        parent1 = "D:/End to End/Evidence"
        path2 = os.path.join(parent1, dirr)
        shutil.rmtree(path2)
            #mylist.clear()
        

    def alter(self):
        mysqldbconnection=dbconnection()
        connection_string=mysqldbconnection.dbconnect_string('conn_string')
        conn=db.connect(**connection_string)
        cursor=conn.cursor()
        query_ip=f"SELECT * FROM py_automate where Run_Indicator='Y';"
        cursor.execute(query_ip)
        column_names = [description[0] for description in cursor.description]
        column_dict = {}
        for column in column_names:
            connection_string=mysqldbconnection.dbconnect_string('conn_string')
            conn=db.connect(**connection_string)
            cursor=conn.cursor()
            cursor.execute("SELECT {} FROM py_automate where Run_Indicator='Y';".format(column))
            column_values = [row[0] for row in cursor.fetchall()]
            column_dict[column] = column_values
        cursor.close()
        conn.close()
        return column_dict
    
    def Screenshot(self,SS_msg):
        mylist1.append(SS_msg)
        directory = "temp"
        parent_dir = "D:\End to End\Evidence"
        path = os.path.join(parent_dir,directory)
        if not os.path.exists(path):
            os.mkdir(path)
        filename = 'D:\End to End\Evidence\\'+directory+'\\'+SS_msg+'.png'
        screen_width, screen_height = pyautogui.size()
        region = (0, 150, screen_width, screen_height - 210)
        screenshot = pyautogui.screenshot(region=region)
        screenshot.save(filename)

    
    def word_cov_alter(self,scn):
        dirr= "temp"
        doc = docx.Document()
        doc.add_heading('Evidence', 0)
        dummy=len(mylist)
        for i in range(dummy):
            doc.add_heading(mylist1[i], 5)
            doc.add_picture('D:\End to End\Evidence\\'+dirr+'\\'+mylist1[i]+'.png', width=Inches(6.5), height=Inches(3.5))
        doc.save('D:\End to End\Evidence\\'+scn+'.docx')
        parent1 = "D:/End to End/Evidence"
        path2 = os.path.join(parent1, dirr)
        shutil.rmtree(path2)
        mylist1.clear()

    def currentdir(self):
        cwd=os.getcwd()
        return cwd
        
        
            
            
        
        


        #res = cursor.fetchall()
    

#sum=py_db_robo()
#sum.DB_Update("PASS","Scenario_01")
#print(mylist)

# l=py_db_robo()
# #print(h.alter())
# print(l.currentdir())



"""
def Db_Connection():
        conn=db.connect(host='localhost',user='root',password='Muruga@2000')
        if conn.is_connected():
            print("DB Connected...")
        else:
            print("connection not established...")
        cursor=conn.cursor()
        cursor.execute("SELECT * FROM pydatabase.py_automate where Run_Indicator='Y';")
        record = cursor.fetchall()
        Scenario,F_Name,L_Name,Email,Comment = ([] for i in range(5))
        for i in range (len(record)):
            Scenario.append(record[i][0])
            F_Name.append(record[i][2])
            L_Name.append(record[i][3])
            Email.append(record[i][4])
            Comment.append(record[i][5])
        mydict={"SCENARIO":Scenario,"F_NAME":F_Name,"L_NAME":L_Name,"EMAIL":Email,"COMMENT":Comment}
        cursor.close()
        conn.close()
        return mydict
"""